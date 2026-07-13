"""基于仓库参考模板生成可审阅但不虚构结论的认证资料副本。"""

from __future__ import annotations

import shutil
from datetime import date
from pathlib import Path

from docx import Document
from openpyxl import load_workbook


ROOT = Path(__file__).resolve().parents[1]
REFERENCE = ROOT / "docs" / "reference"
OUTPUT = ROOT / "docs" / "selftest" / "certification"
TODAY = date(2026, 7, 13).isoformat()


def first_existing(*names: str) -> Path:
    for name in names:
        candidate = REFERENCE / name
        if candidate.exists():
            return candidate
    raise FileNotFoundError(f"参考模板不存在: {names}")


def fill_workbook() -> Path:
    source = first_existing("XXX-BMS 软件静态分析报告.xlsx")
    target = OUTPUT / "TLSR8258-BMS_软件静态分析报告_已填写.xlsx"
    shutil.copy2(source, target)
    workbook = load_workbook(target)

    replacements = {
        "XXX-BMS": "TLSR8258-SH367309 BMS",
        "XXX BMS": "TLSR8258-SH367309 BMS",
        "XXX": "TLSR8258-SH367309",
    }
    for sheet in workbook.worksheets:
        for row in sheet.iter_rows():
            for cell in row:
                if isinstance(cell.value, str):
                    value = cell.value
                    for old, new in replacements.items():
                        value = value.replace(old, new)
                    cell.value = value

    cover = workbook["Cover"]
    cover["A1"] = "TLSR8258-SH367309 BMS 软件静态分析报告"
    cover["A2"] = "状态：待执行独立静态分析；本文件不表示 MISRA/认证通过"
    cover["A3"] = f"模板生成日期：{TODAY}；执行人/复核人：待填写"

    history = workbook["ChangeHistory"]
    history.append(["0.1", TODAY, "根据真实工程生成项目化模板；测试结论保持待执行", "Codex", "待复核"])

    general = workbook["General"]
    general.append(["Project", "TLSR8258-SH367309 BMS / tc_ble_single_sdk V3.4.2.8 Patch 0001"])
    general.append(["Compiler", r"TC32 GCC: C:\TelinkSDK\opt\tc32\bin"])
    general.append(["Scope", "vendor/ble_sample/bms_selftest 及其启动、链接、AFE、Modbus 接入"])
    general.append(["Build evidence", "2026-07-13 clean build 成功；新增模块无新告警；原工程告警未据此判为合格"])
    general.append(["Static analysis", "待使用项目批准版本的 Cppcheck/MISRA 工具执行并人工复核"])
    general.append(["Certification boundary", "仅为认证准备证据模板，不构成 IEC 60335/IEC 60730 合规结论"])

    if "DevitionList" in workbook.sheetnames:
        deviation = workbook["DevitionList"]
        deviation.append(["DEV-TC32-001", "TC32 专有汇编/编译器扩展", "MISRA C 工具不分析汇编", "反汇编审查+目标故障注入", "待批准"])
        deviation.append(["DEV-HW-001", "缺少独立时钟参考", "仅检测停止/严重调度异常", "增加外部参考或安全手册论证", "待处理"])

    workbook.save(target)
    return target


def fill_validation_spec() -> Path:
    source = first_existing("XXX-BMS 软件验证规范.docx")
    target = OUTPUT / "TLSR8258-BMS_软件验证规范_已填写.docx"
    shutil.copy2(source, target)
    document = Document(target)

    for paragraph in document.paragraphs:
        if paragraph.text:
            for run in paragraph.runs:
                run.text = run.text.replace("XXX-BMS", "TLSR8258-SH367309 BMS").replace(
                    "XXX BMS", "TLSR8258-SH367309 BMS"
                )

    document.add_page_break()
    document.add_heading("附录：MCU 自检与安全降级验证（项目化补充）", level=1)
    document.add_paragraph(
        "生成日期：" + TODAY + "。本附录引用 docs/selftest 下的追踪矩阵、测试计划和残余风险。"
        "“未执行”项目不得在无目标板原始证据时改为通过。"
    )
    table = document.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    for index, title in enumerate(("编号", "测试", "故障模型", "检测机制", "预期安全态", "当前状态")):
        table.rows[0].cells[index].text = title
    cases = (
        ("SW-VD-MCU-001", "CPU 寄存器/标志", "寄存器/执行单元固定故障", "TC32 汇编模式与分支", "所有 MOS/均衡/CTLC/RF_EN 关闭", "代码/构建完成，板测未执行"),
        ("SW-VD-MCU-002", "控制流", "非法跳转/计数损坏", "签名及正反码", "致命锁存、停止喂狗", "代码/构建完成，板测未执行"),
        ("SW-VD-MCU-003", "Flash", "程序位翻转", "Manifest CRC + Telink CRC", "启动前禁止输出", "主机测试通过，板测未执行"),
        ("SW-VD-MCU-004", "RAM/栈", "RAM 固定位/越界", "专用区 March-like + guard", "致命锁存", "布局检查通过，板测未执行"),
        ("SW-VD-MCU-005", "时钟/中断", "停止/严重异常", "tick 与 Timer0 窗口", "安全态/看门狗复位", "代码完成，独立频率覆盖不足"),
        ("SW-VD-BMS-001", "AFE/ADC/MOS", "通信、配置、反馈异常", "回读/合理性/一致性", "关闭功率与均衡输出", "代码完成，夹具测试未执行"),
        ("SW-VD-SYS-001", "OTA/retention", "镜像损坏/唤醒异常", "双 CRC/唤醒重检", "不进入正常输出", "OTA 主机校验通过；retention 当前关闭"),
    )
    for case in cases:
        cells = table.add_row().cells
        for index, value in enumerate(case):
            cells[index].text = value

    document.add_heading("证据索引", level=2)
    for item in (
        "docs/selftest/02_requirements_traceability.md",
        "docs/selftest/03_certification_test_plan.md",
        "docs/selftest/04_fault_injection_matrix.md",
        "docs/selftest/06_test_report_template.md",
        "docs/selftest/07_residual_risk.md",
        "project/tlsr_tc32/B85/825x_ble_sample/selftest_image_report.json",
        "project/tlsr_tc32/B85/825x_ble_sample/selftest_layout_report.json",
    ):
        document.add_paragraph("• " + item)
    document.add_paragraph("执行人、复核人、硬件版本、工具精确版本、波形和实际检测时间：待填写。")
    document.save(target)
    return target


def main() -> None:
    OUTPUT.mkdir(parents=True, exist_ok=True)
    workbook = fill_workbook()
    document = fill_validation_spec()
    print(workbook)
    print(document)


if __name__ == "__main__":
    main()
